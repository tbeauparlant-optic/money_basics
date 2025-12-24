import re
import difflib
from pathlib import Path
from urllib.parse import quote
from docx import Document

# -----------------------------
# CONFIG — EDIT THESE 3 LINES
# -----------------------------
PART_DOCX_DIR = Path("Parts")          # where your Part*.docx live (change if different)
MANIFEST_PATH = Path("repo_manifest.txt")
OUTPUT_DIR = Path("Parts_with_repo_snippets")
BASE_URL = "{{BASE_URL}}"             # later replace once (e.g., GitHub Pages base)

# If you want to only target Workshop 6 assets:
WORKSHOP_FILTER = re.compile(r"Parts/Teaching Materials/Workshop 6/", re.IGNORECASE)

KW_RE = re.compile(r"\b(handout|handouts|worksheet|worksheets)\b", re.IGNORECASE)

def url_for_repo_path(repo_rel_path: str) -> str:
    # URL-encode each path segment, preserve slashes
    parts = [quote(seg) for seg in repo_rel_path.split("/")]
    return f"{BASE_URL}/" + "/".join(parts)

def norm(s: str) -> str:
    s = s.lower()
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def load_manifest():
    paths = []
    for line in MANIFEST_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line:
            continue
        if not line.lower().endswith(".pdf"):
            continue
        # Optional: only Workshop 6 handouts/resources
        if WORKSHOP_FILTER and not WORKSHOP_FILTER.search(line):
            continue
        paths.append(line)
    return paths

def best_match(reference_text: str, pdf_paths: list[str]):
    ref = norm(reference_text)

    # candidates based on filename stem
    cand = [(p, norm(Path(p).stem)) for p in pdf_paths]

    # direct containment wins
    for p, stem in cand:
        if stem and stem in ref:
            return p, 1.0

    # fuzzy match against stems
    stems = [stem for _, stem in cand]
    best = difflib.get_close_matches(ref, stems, n=1, cutoff=0.55)
    if best:
        best_stem = best[0]
        for p, stem in cand:
            if stem == best_stem:
                score = difflib.SequenceMatcher(None, ref, stem).ratio()
                return p, score

    return None, 0.0

def insert_after(doc: Document, para, text_block: str):
    new_p = doc.add_paragraph(text_block)
    para._p.addnext(new_p._p)

def process_docx(docx_path: Path, pdf_paths: list[str]):
    doc = Document(docx_path)
    changed = False

    for para in list(doc.paragraphs):
        txt = para.text.strip()
        if not txt:
            continue
        if not KW_RE.search(txt):
            continue

        ref_type = "Handout" if "handout" in txt.lower() else "Worksheet"
        match, score = best_match(txt, pdf_paths)

        if match:
            url = url_for_repo_path(match)
            snippet = "\n".join([
                "[MB_ASSET]",
                f"Type: {ref_type}",
                f"Reference: {txt}",
                f"Repo path: {match}",
                f"URL: {url}",
                "[/MB_ASSET]",
            ])
        else:
            snippet = "\n".join([
                "[MB_ASSET]",
                f"Type: {ref_type}",
                f"Reference: {txt}",
                "Repo path: NOT FOUND",
                "URL: ",
                "[/MB_ASSET]",
            ])

        insert_after(doc, para, snippet)
        changed = True

    return doc, changed

def main():
    if not MANIFEST_PATH.exists():
        raise SystemExit(f"Missing {MANIFEST_PATH}. Run: git ls-files ... > repo_manifest.txt")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf_paths = load_manifest()
    if not pdf_paths:
        raise SystemExit("Manifest loaded, but no PDFs matched the filter. Check WORKSHOP_FILTER or manifest content.")

    docx_files = sorted(PART_DOCX_DIR.rglob("Part*.docx"))
    if not docx_files:
        raise SystemExit(f"No Part*.docx found under {PART_DOCX_DIR}")

    for f in docx_files:
        doc, changed = process_docx(f, pdf_paths)
        out = OUTPUT_DIR / f.name
        doc.save(out)
        print(("UPDATED" if changed else "NOREFS"), f"->", out)

    print("Done:", OUTPUT_DIR)

if __name__ == "__main__":
    main()
